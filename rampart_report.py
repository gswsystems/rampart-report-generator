#!/usr/bin/env python3
"""
Rampart Report Generator

Generates Word documents from Rampart JSON exports using custom .docx templates.
Templates use Jinja2-style placeholders ({{ variable }}) for text and special
table markers for repeating data.

Usage:
    python rampart_report.py input.json template.docx output.docx
    python rampart_report.py input.json template.docx output.docx --date "2026-03-14"
    python rampart_report.py --list-variables input.json
"""

import argparse
import json
import sys
import re
import os
from datetime import datetime
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from jinja2 import Environment


def load_json(path):
    """Load and return the Rampart JSON export."""
    with open(path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Handle full report format (has 'analysis' wrapper) vs direct export
    if 'analysis' in data and 'configuration' in data:
        return data['analysis'], data.get('configuration')
    return data, None


def build_variables(analysis, config=None, overrides=None):
    """
    Build the flat variable dictionary from analysis JSON.
    Returns a dict of all available template variables.
    """
    v = {}
    overrides = overrides or {}

    summary = analysis.get('summary', {})
    risk = analysis.get('risk_rating', {})
    severity = summary.get('severity_breakdown', {})

    # --- Report metadata ---
    v['report_date'] = overrides.get('date', datetime.now().strftime('%Y-%m-%d'))
    v['report_title'] = overrides.get('title', 'Firewall Security Audit Report')
    v['client_name'] = overrides.get('client_name', '')
    v['client_contact'] = overrides.get('client_contact', '')
    v['auditor_name'] = overrides.get('auditor_name', '')
    v['auditor_company'] = overrides.get('auditor_company', '')
    v['confidentiality'] = overrides.get('confidentiality', 'CONFIDENTIAL')

    # --- Summary ---
    v['total_rules'] = summary.get('total_rules_analyzed', 0)
    v['rules_with_issues'] = summary.get('rules_with_issues', 0)
    v['compliance_rate'] = round(summary.get('compliance_rate', 0), 1)
    v['config_type'] = summary.get('config_type', '')
    v['analysis_timestamp'] = summary.get('timestamp', '')
    v['device_group_count'] = summary.get('device_group_count', 0)
    v['device_groups'] = ', '.join(summary.get('device_groups', []))

    # --- Severity counts ---
    v['critical_count'] = severity.get('Critical', 0)
    v['high_count'] = severity.get('High', 0)
    v['medium_count'] = severity.get('Medium', 0)
    v['low_count'] = severity.get('Low', 0)
    v['total_findings'] = v['critical_count'] + v['high_count'] + v['medium_count'] + v['low_count']

    # --- Risk rating ---
    v['risk_score'] = risk.get('score', 0)
    v['risk_grade'] = risk.get('grade', 'N/A')
    v['best_practices_score'] = risk.get('best_practices_score', 0)
    v['segmentation_score'] = risk.get('segmentation_score', 0)
    v['critical_issues'] = risk.get('critical_issues', 0)
    v['high_risk_rules'] = risk.get('high_risk_rules', 0)
    v['shadowed_rule_count'] = risk.get('shadowed_rules', 0)
    v['lateral_movement_paths'] = risk.get('lateral_movement_paths', 0)

    # --- Duplicate objects ---
    dupes = analysis.get('duplicate_objects', {})
    v['duplicate_address_count'] = dupes.get('total_duplicate_addresses', 0)
    v['duplicate_service_count'] = dupes.get('total_duplicate_services', 0)

    # --- Shadowed rules ---
    shadowed = analysis.get('shadowed_rules', [])
    v['shadowed_rules_total'] = len(shadowed)

    # --- Compliance ---
    compliance = analysis.get('compliance', {})
    v['compliance_data_available'] = bool(compliance)

    # --- Attack surface ---
    attack = analysis.get('attack_surface', {})
    v['attack_surface_available'] = bool(attack)

    # --- Best practices ---
    bp = analysis.get('best_practices', {})
    v['best_practices_available'] = bool(bp)
    if bp:
        v['best_practices_overall_score'] = bp.get('overall_score', 0)
        v['best_practices_grade'] = bp.get('grade', 'N/A')

    # --- Analyzer result counts ---
    def count_findings(key, sub_key=None):
        data = analysis.get(key, {})
        if sub_key:
            return len(data.get(sub_key, []))
        if isinstance(data, list):
            return len(data)
        return data.get('total_issues', data.get('total_findings', 0))

    v['rule_expiry_count'] = (
        analysis.get('rule_expiry', {}).get('expired_schedule_count', 0) +
        analysis.get('rule_expiry', {}).get('likely_temporary_count', 0)
    )
    v['cleartext_rule_count'] = analysis.get('cleartext_exposure', {}).get('cleartext_rule_count', 0)
    v['geo_ip_unrestricted_count'] = analysis.get('geo_ip_exposure', {}).get('unrestricted_count', 0)
    v['lateral_movement_count'] = analysis.get('lateral_movement', {}).get('total_issues', 0)
    v['stale_rule_count'] = (
        analysis.get('stale_rules', {}).get('stale_named_count', 0) +
        analysis.get('stale_rules', {}).get('unused_object_rule_count', 0)
    )
    v['egress_risk_count'] = len(analysis.get('egress_filtering', {}).get('findings', []))
    v['decryption_gap_count'] = analysis.get('decryption_policy', {}).get('gaps_count', 0)

    # --- Segmentation ---
    seg = analysis.get('segmentation_score', {}).get('score', {})
    v['seg_score'] = seg.get('segmentation_score', 0)
    v['seg_grade'] = seg.get('grade', 'N/A')
    v['seg_zone_count'] = seg.get('zone_count', 0)
    v['seg_allowed_pairs'] = seg.get('allowed_pairs', 0)
    v['seg_blocked_pairs'] = seg.get('blocked_pairs', 0)

    return v


def build_table_data(analysis):
    """
    Build table datasets that can be inserted into template tables.
    Returns a dict of table_name -> list of row dicts.
    """
    tables = {}

    # --- Findings table ---
    findings_rows = []
    for rf in analysis.get('findings', []):
        for f in rf.get('findings', []):
            findings_rows.append({
                'rule_name': rf.get('rule_name', ''),
                'device_group': rf.get('device_group', ''),
                'severity': f.get('severity', ''),
                'type': f.get('type', ''),
                'description': f.get('description', ''),
                'remediation': f.get('remediation', ''),
                'risk_score': rf.get('risk_score', 0),
            })
    tables['findings'] = findings_rows

    # --- Critical/High findings only ---
    tables['critical_findings'] = [r for r in findings_rows if r['severity'] == 'Critical']
    tables['high_findings'] = [r for r in findings_rows if r['severity'] == 'High']

    # --- Shadowed rules ---
    tables['shadowed_rules'] = [
        {
            'rule_name': s.get('shadowed_rule_name', ''),
            'shadowed_by': s.get('shadowed_by_rule_name', ''),
            'device_group': s.get('device_group', ''),
            'severity': s.get('severity', ''),
            'description': s.get('description', ''),
            'remediation': s.get('remediation', ''),
        }
        for s in analysis.get('shadowed_rules', [])
    ]

    # --- Duplicate addresses ---
    tables['duplicate_addresses'] = [
        {
            'type': d.get('type', ''),
            'value': d.get('value', ''),
            'count': d.get('duplicate_count', 0),
            'objects': ', '.join(d.get('object_names', [])),
            'remediation': d.get('remediation', ''),
        }
        for d in analysis.get('duplicate_objects', {}).get('duplicate_addresses', [])
    ]

    # --- Compliance frameworks ---
    compliance = analysis.get('compliance', {})
    if isinstance(compliance, dict):
        frameworks = []
        for key, val in compliance.items():
            if isinstance(val, dict) and 'compliance_percentage' in val:
                frameworks.append({
                    'framework': key,
                    'percentage': val.get('compliance_percentage', 0),
                    'status': val.get('status', ''),
                    'passed': val.get('passed_controls', 0),
                    'failed': val.get('failed_controls', 0),
                    'total': val.get('total_controls', 0),
                })
        tables['compliance'] = frameworks

    # --- Lateral movement ---
    tables['lateral_movement'] = [
        {
            'rule_name': f.get('rule_name', ''),
            'severity': f.get('severity', ''),
            'source_zones': ', '.join(f.get('source_zones', [])),
            'dest_zones': ', '.join(f.get('destination_zones', [])),
            'risk_factors': '; '.join(f.get('risk_factors', [])),
        }
        for f in analysis.get('lateral_movement', {}).get('lateral_movement_rules', [])
    ]

    # --- Weak segments ---
    tables['weak_segments'] = [
        {
            'source_zone': s.get('source_zone', ''),
            'dest_zone': s.get('destination_zone', ''),
            'openness': s.get('openness', ''),
            'remediation': s.get('remediation', ''),
        }
        for s in analysis.get('segmentation_score', {}).get('weak_segments', [])
    ]

    # --- Egress findings ---
    tables['egress_findings'] = [
        {
            'rule_name': f.get('rule_name', ''),
            'severity': f.get('severity', ''),
            'risk_factors': '; '.join(f.get('risk_factors', [])),
            'remediation': f.get('remediation', ''),
        }
        for f in analysis.get('egress_filtering', {}).get('findings', [])
    ]

    # --- Cleartext exposure ---
    tables['cleartext_rules'] = [
        {
            'rule_name': f.get('rule_name', ''),
            'protocol': f.get('protocol', ''),
            'severity': f.get('severity', ''),
            'secure_alternative': f.get('secure_alternative', ''),
        }
        for f in analysis.get('cleartext_exposure', {}).get('cleartext_rules', [])
    ]

    # --- Stale rules ---
    stale = analysis.get('stale_rules', {})
    stale_rows = []
    for r in stale.get('stale_named_rules', []):
        stale_rows.append({
            'rule_name': r.get('rule_name', ''),
            'severity': r.get('severity', ''),
            'indicators': '; '.join(r.get('indicators', [])),
            'disabled': 'Yes' if r.get('disabled') else 'No',
        })
    for r in stale.get('unused_object_rules', []):
        stale_rows.append({
            'rule_name': r.get('rule_name', ''),
            'severity': 'High',
            'indicators': 'References missing objects',
            'disabled': 'No',
        })
    tables['stale_rules'] = stale_rows

    # --- Decryption gaps ---
    tables['decryption_gaps'] = [
        {
            'rule_name': g.get('rule_name', ''),
            'severity': g.get('severity', ''),
            'reason': g.get('reason', ''),
            'remediation': g.get('remediation', ''),
        }
        for g in analysis.get('decryption_policy', {}).get('gaps', [])
    ]

    # --- Geo-IP exposure ---
    geo = analysis.get('geo_ip_exposure', {})
    geo_rows = []
    for r in geo.get('unrestricted_external_rules', []):
        geo_rows.append({
            'rule_name': r.get('rule_name', ''),
            'severity': r.get('severity', ''),
            'type': 'Unrestricted External',
            'remediation': r.get('remediation', ''),
        })
    for r in geo.get('missing_geo_block_rules', []):
        geo_rows.append({
            'rule_name': r.get('rule_name', ''),
            'severity': r.get('severity', ''),
            'type': 'Missing Geo-Block',
            'remediation': r.get('remediation', ''),
        })
    tables['geo_ip_findings'] = geo_rows

    # --- Rule expiry ---
    expiry = analysis.get('rule_expiry', {})
    expiry_rows = []
    for r in expiry.get('expired_schedule_rules', []):
        expiry_rows.append({
            'rule_name': r.get('rule_name', ''),
            'type': 'Expired Schedule',
            'detail': f"Schedule: {r.get('schedule', '')}, expired {r.get('days_expired', 0)} days ago",
        })
    for r in expiry.get('likely_temporary_rules', []):
        expiry_rows.append({
            'rule_name': r.get('rule_name', ''),
            'type': 'Likely Temporary',
            'detail': r.get('reason', ''),
        })
    tables['rule_expiry'] = expiry_rows

    return tables


# ---------------------------------------------------------------------------
# Template processing
# ---------------------------------------------------------------------------

# Regex to match {{ variable }} placeholders in text
PLACEHOLDER_RE = re.compile(r'\{\{\s*(\w+)\s*\}\}')

# Regex to match table row markers: {{#table_name}} ... {{/table_name}}
TABLE_START_RE = re.compile(r'\{\{#(\w+)\}\}')
TABLE_END_RE = re.compile(r'\{\{/(\w+)\}\}')


def get_severity_color(severity):
    """Return an RGBColor for the given severity string."""
    s = severity.lower() if severity else ''
    if s == 'critical':
        return RGBColor(0xDC, 0x26, 0x26)
    elif s == 'high':
        return RGBColor(0xEF, 0x44, 0x44)
    elif s == 'medium':
        return RGBColor(0xF5, 0x9E, 0x0B)
    elif s == 'low':
        return RGBColor(0x64, 0x74, 0x8B)
    return None


def replace_in_run(run, variables):
    """Replace {{ var }} placeholders in a single run's text."""
    if not run.text:
        return
    text = run.text

    def replacer(match):
        key = match.group(1)
        val = variables.get(key)
        if val is None:
            return match.group(0)  # Leave unknown placeholders as-is
        return str(val)

    run.text = PLACEHOLDER_RE.sub(replacer, text)


def replace_in_paragraph(paragraph, variables):
    """
    Replace placeholders in a paragraph. Handles placeholders split
    across multiple runs by first joining, replacing, then re-splitting.
    """
    full_text = ''.join(r.text or '' for r in paragraph.runs)
    if '{{' not in full_text:
        return

    # If placeholder is contained within a single run, replace directly
    for run in paragraph.runs:
        if run.text and '{{' in run.text and '}}' in run.text:
            replace_in_run(run, variables)

    # Check if there are still unresolved placeholders split across runs
    full_text = ''.join(r.text or '' for r in paragraph.runs)
    if '{{' not in full_text:
        return

    # Merge all runs, replace, put text in first run, clear rest
    new_text = PLACEHOLDER_RE.sub(
        lambda m: str(variables.get(m.group(1), m.group(0))),
        full_text
    )
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''


def process_table(table, variables, table_data):
    """
    Process a table in the document.

    If a row contains {{#table_name}}, it marks a template row.
    That row is cloned for each item in the corresponding table_data list,
    with {{column}} placeholders replaced by row values.
    The marker row is then removed.

    Otherwise, simple {{variable}} placeholders in cells are replaced.
    """
    rows_to_remove = []
    rows_to_insert = []

    for row_idx, row in enumerate(table.rows):
        row_text = ''.join(c.text for c in row.cells)

        # Check for table row marker
        start_match = TABLE_START_RE.search(row_text)
        if start_match:
            table_name = start_match.group(1)
            data_rows = table_data.get(table_name, [])

            if not data_rows:
                rows_to_remove.append(row_idx)
                continue

            rows_to_remove.append(row_idx)

            # Use this row as the template
            template_row = row
            for data_idx, data_row in enumerate(data_rows):
                rows_to_insert.append((row_idx + data_idx, template_row, data_row, table_name))
            continue

        # Simple placeholder replacement in non-template rows
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, variables)

    # Insert data rows (in reverse to maintain indices)
    for insert_idx, (target_idx, template_row, data_row, table_name) in enumerate(rows_to_insert):
        new_row = deepcopy(template_row._tr)
        # Replace placeholders in the new row
        from docx.oxml.ns import qn
        for tc in new_row.findall(qn('w:tc')):
            for p in tc.findall(qn('w:p')):
                full = ''.join(
                    r.text or '' for r in p.findall(qn('w:r'))
                )
                # Remove the {{#name}} and {{/name}} markers
                full = TABLE_START_RE.sub('', full)
                full = TABLE_END_RE.sub('', full)
                # Replace column placeholders
                new_text = PLACEHOLDER_RE.sub(
                    lambda m: str(data_row.get(m.group(1), m.group(0))),
                    full
                )
                runs = p.findall(qn('w:r'))
                if runs:
                    # Put all text in first run, clear rest
                    t_elem = runs[0].find(qn('w:t'))
                    if t_elem is not None:
                        t_elem.text = new_text
                        t_elem.set(qn('xml:space'), 'preserve')
                    for r in runs[1:]:
                        t_elem = r.find(qn('w:t'))
                        if t_elem is not None:
                            t_elem.text = ''

        table._tbl.append(new_row)

    # Remove marker rows (in reverse order)
    for row_idx in sorted(rows_to_remove, reverse=True):
        row_elem = table.rows[row_idx]._tr
        table._tbl.remove(row_elem)


def process_document(doc, variables, table_data):
    """Process all paragraphs and tables in the document."""
    # Process paragraphs in body
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, variables)

    # Process tables
    for table in doc.tables:
        process_table(table, variables, table_data)

    # Process headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header and header.is_linked_to_previous is False:
                for paragraph in header.paragraphs:
                    replace_in_paragraph(paragraph, variables)
                for table in header.tables:
                    process_table(table, variables, table_data)
        for footer in [section.footer, section.first_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                for paragraph in footer.paragraphs:
                    replace_in_paragraph(paragraph, variables)
                for table in footer.tables:
                    process_table(table, variables, table_data)


def list_variables(analysis):
    """Print all available template variables and their current values."""
    variables = build_variables(analysis)
    table_data = build_table_data(analysis)

    print("=" * 60)
    print("TEMPLATE VARIABLES")
    print("=" * 60)
    print()
    print("Use these in your template as {{ variable_name }}")
    print()

    # Group by category
    groups = {
        'Report Metadata': ['report_date', 'report_title', 'client_name', 'client_contact',
                           'auditor_name', 'auditor_company', 'confidentiality'],
        'Summary': ['total_rules', 'rules_with_issues', 'compliance_rate', 'config_type',
                    'analysis_timestamp', 'device_group_count', 'device_groups'],
        'Severity Counts': ['critical_count', 'high_count', 'medium_count', 'low_count',
                           'total_findings'],
        'Risk Rating': ['risk_score', 'risk_grade', 'best_practices_score', 'segmentation_score',
                       'critical_issues', 'high_risk_rules', 'shadowed_rule_count',
                       'lateral_movement_paths'],
        'Duplicate Objects': ['duplicate_address_count', 'duplicate_service_count'],
        'Shadowed Rules': ['shadowed_rules_total'],
        'Best Practices': ['best_practices_available', 'best_practices_overall_score',
                          'best_practices_grade'],
        'Segmentation': ['seg_score', 'seg_grade', 'seg_zone_count', 'seg_allowed_pairs',
                         'seg_blocked_pairs'],
        'Analyzer Counts': ['rule_expiry_count', 'cleartext_rule_count',
                           'geo_ip_unrestricted_count', 'lateral_movement_count',
                           'stale_rule_count', 'egress_risk_count', 'decryption_gap_count'],
    }

    for group_name, keys in groups.items():
        print(f"  {group_name}:")
        for key in keys:
            val = variables.get(key, '')
            print(f"    {{ {key} }} = {val}")
        print()

    print("=" * 60)
    print("TABLE DATA")
    print("=" * 60)
    print()
    print("Use these in table rows with {{#table_name}} ... {{/table_name}}")
    print()

    for name, rows in table_data.items():
        if rows:
            cols = list(rows[0].keys())
            print(f"  {{{{#{name}}}}} ({len(rows)} rows)")
            print(f"    Columns: {', '.join(cols)}")
        else:
            print(f"  {{{{#{name}}}}} (0 rows)")
        print()


def generate_report(json_path, template_path, output_path, overrides=None):
    """Main report generation function."""
    analysis, config = load_json(json_path)
    variables = build_variables(analysis, config, overrides)
    table_data = build_table_data(analysis)

    doc = Document(template_path)
    process_document(doc, variables, table_data)
    doc.save(output_path)

    return output_path


def main():
    parser = argparse.ArgumentParser(
        description='Generate Word documents from Rampart JSON exports using custom templates.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s audit.json template.docx report.docx
  %(prog)s audit.json template.docx report.docx --client "Acme Corp"
  %(prog)s --list-variables audit.json
        """
    )

    parser.add_argument('json_file', help='Rampart JSON export file')
    parser.add_argument('template', nargs='?', help='Word template (.docx)')
    parser.add_argument('output', nargs='?', help='Output file path (.docx)')

    parser.add_argument('--list-variables', action='store_true',
                       help='List all available template variables and exit')
    parser.add_argument('--client', dest='client_name', default='',
                       help='Client name for the report')
    parser.add_argument('--client-contact', default='',
                       help='Client contact person')
    parser.add_argument('--auditor', dest='auditor_name', default='',
                       help='Auditor name')
    parser.add_argument('--company', dest='auditor_company', default='',
                       help='Auditing company name')
    parser.add_argument('--title', default='Firewall Security Audit Report',
                       help='Report title')
    parser.add_argument('--date', default=None,
                       help='Report date (default: today)')
    parser.add_argument('--confidentiality', default='CONFIDENTIAL',
                       help='Confidentiality level')

    args = parser.parse_args()

    # Load JSON
    try:
        analysis, config = load_json(args.json_file)
    except FileNotFoundError:
        print(f"Error: JSON file not found: {args.json_file}", file=sys.stderr)
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON: {e}", file=sys.stderr)
        sys.exit(1)

    if args.list_variables:
        list_variables(analysis)
        sys.exit(0)

    if not args.template or not args.output:
        parser.error('template and output arguments are required (unless using --list-variables)')

    if not os.path.exists(args.template):
        print(f"Error: Template file not found: {args.template}", file=sys.stderr)
        sys.exit(1)

    overrides = {
        'client_name': args.client_name,
        'client_contact': args.client_contact,
        'auditor_name': args.auditor_name,
        'auditor_company': args.auditor_company,
        'title': args.title,
        'confidentiality': args.confidentiality,
    }
    if args.date:
        overrides['date'] = args.date

    try:
        output = generate_report(args.json_file, args.template, args.output, overrides)
        print(f"Report generated: {output}")
    except Exception as e:
        print(f"Error generating report: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
