#!/usr/bin/env python3
"""
Creates a starter Word template for the Rampart Report Generator.

Run this script to generate 'template.docx' which contains all available
placeholders, table markers, and professional formatting as a starting point.

Usage:
    python create_template.py [output_path]
"""

import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=Pt(10), color=None, alignment=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = size
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color


def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    return h


def add_table_with_headers(doc, headers, col_widths=None):
    """Create a formatted table with header row."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "1B3A5C")
        set_cell_text(cell, header, bold=True, size=Pt(9),
                      color=RGBColor(0xFF, 0xFF, 0xFF))

    if col_widths:
        for i, width in enumerate(col_widths):
            table.columns[i].width = width

    return table


def add_table_marker_row(table, marker_name, placeholders):
    """Add start marker, data row, and end marker row to a table."""
    # Start marker row
    start_row = table.add_row()
    start_row.cells[0].text = "{{#" + marker_name + "}}"
    for i in range(1, len(start_row.cells)):
        start_row.cells[i].text = ""

    # Data row with placeholders
    data_row = table.add_row()
    for i, ph in enumerate(placeholders):
        cell = data_row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run("{{ " + ph + " }}")
        run.font.size = Pt(9)
        run.font.name = "Calibri"

    # End marker row
    end_row = table.add_row()
    end_row.cells[0].text = "{{/" + marker_name + "}}"
    for i in range(1, len(end_row.cells)):
        end_row.cells[i].text = ""


def create_template(output_path):
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)  # A4
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # -- Header --
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("{{ auditor_company }}  |  {{ confidentiality }}")
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    hr.font.name = "Calibri"

    # -- Footer --
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("{{ report_title }}  |  {{ client_name }}  |  {{ report_date }}")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    fr.font.name = "Calibri"

    # =========================================================================
    # COVER PAGE
    # =========================================================================
    for _ in range(6):
        doc.add_paragraph("")

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("{{ report_title }}")
    title_run.font.size = Pt(32)
    title_run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    title_run.bold = True
    title_run.font.name = "Calibri"

    # Subtitle line
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Prepared for {{ client_name }}")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    sub_run.font.name = "Calibri"

    doc.add_paragraph("")

    # Metadata table on cover
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_items = [
        ("Date", "{{ report_date }}"),
        ("Client Contact", "{{ client_contact }}"),
        ("Auditor", "{{ auditor_name }}"),
        ("Company", "{{ auditor_company }}"),
        ("Confidentiality", "{{ confidentiality }}"),
        ("Document Version", "1.0"),
    ]
    for i, (label, value) in enumerate(meta_items):
        set_cell_shading(meta_table.rows[i].cells[0], "F2F2F2")
        set_cell_text(meta_table.rows[i].cells[0], label, bold=True, size=Pt(10))
        set_cell_text(meta_table.rows[i].cells[1], value, size=Pt(10))
    meta_table.columns[0].width = Inches(2.0)
    meta_table.columns[1].width = Inches(3.5)

    # Page break after cover
    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS (placeholder)
    # =========================================================================
    add_heading_styled(doc, "Table of Contents", level=1)

    # Insert a Word TOC field that generates from Heading 1-3 styles.
    # The TOC will show placeholder text until the user right-clicks it
    # in Word and selects "Update Field" (or presses F9).
    toc_p = doc.add_paragraph()
    run = toc_p.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fld_char_begin)

    run2 = toc_p.add_run()
    instr = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run2._r.append(instr)

    run3 = toc_p.add_run()
    fld_char_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fld_char_separate)

    # Placeholder text shown before the field is updated
    run4 = toc_p.add_run("Right-click and select 'Update Field' to generate the table of contents.")
    run4.font.size = Pt(10)
    run4.font.italic = True
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    run5 = toc_p.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fld_char_end)

    doc.add_page_break()

    # =========================================================================
    # 1. EXECUTIVE SUMMARY
    # =========================================================================
    add_heading_styled(doc, "1. Executive Summary", level=1)

    doc.add_paragraph(
        "This report presents the findings of a firewall security audit conducted "
        "for {{ client_name }} on {{ report_date }}. The audit analysed "
        "{{ total_rules }} firewall rules across {{ device_group_count }} device "
        "group(s) ({{ device_groups }})."
    )

    doc.add_paragraph("")

    # Summary metrics table
    add_heading_styled(doc, "1.1 Key Metrics", level=2)
    metrics_table = doc.add_table(rows=4, cols=4)
    metrics_table.style = "Table Grid"
    metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    metrics = [
        ("Total Rules", "{{ total_rules }}", "Compliance Rate", "{{ compliance_rate }}"),
        ("Rules with Issues", "{{ rules_with_issues }}", "Risk Score", "{{ risk_score }}"),
        ("Total Findings", "{{ total_findings }}", "Risk Grade", "{{ risk_grade }}"),
        ("Shadowed Rules", "{{ shadowed_rule_count }}", "Best Practices Score", "{{ best_practices_overall_score }}"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(metrics):
        set_cell_shading(metrics_table.rows[i].cells[0], "F2F2F2")
        set_cell_text(metrics_table.rows[i].cells[0], l1, bold=True, size=Pt(9))
        set_cell_text(metrics_table.rows[i].cells[1], v1, size=Pt(9))
        set_cell_shading(metrics_table.rows[i].cells[2], "F2F2F2")
        set_cell_text(metrics_table.rows[i].cells[2], l2, bold=True, size=Pt(9))
        set_cell_text(metrics_table.rows[i].cells[3], v2, size=Pt(9))

    doc.add_paragraph("")

    # Severity breakdown
    add_heading_styled(doc, "1.2 Findings by Severity", level=2)
    sev_table = doc.add_table(rows=5, cols=2)
    sev_table.style = "Table Grid"
    sev_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sev_items = [
        ("Critical", "{{ critical_count }}", "C00000"),
        ("High", "{{ high_count }}", "E36C09"),
        ("Medium", "{{ medium_count }}", "E6B800"),
        ("Low", "{{ low_count }}", "4472C4"),
        ("Total", "{{ total_findings }}", "1B3A5C"),
    ]
    # Header
    set_cell_shading(sev_table.rows[0].cells[0], "1B3A5C")
    set_cell_text(sev_table.rows[0].cells[0], "Severity", bold=True, size=Pt(9),
                  color=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_shading(sev_table.rows[0].cells[1], "1B3A5C")
    set_cell_text(sev_table.rows[0].cells[1], "Count", bold=True, size=Pt(9),
                  color=RGBColor(0xFF, 0xFF, 0xFF))
    # Rearrange: header + 4 severity rows
    sev_table2 = doc.add_table(rows=1, cols=2)
    # Actually, let's just redo this properly
    doc.element.body.remove(sev_table._tbl)
    doc.element.body.remove(sev_table2._tbl)

    sev_table = add_table_with_headers(doc, ["Severity", "Count"])
    for label, placeholder, color in sev_items:
        row = sev_table.add_row()
        set_cell_text(row.cells[0], label, bold=True, size=Pt(9),
                      color=RGBColor.from_string(color))
        set_cell_text(row.cells[1], placeholder, size=Pt(9))
    sev_table.columns[0].width = Inches(2.0)
    sev_table.columns[1].width = Inches(1.5)

    doc.add_paragraph("")

    # =========================================================================
    # 2. RISK ASSESSMENT
    # =========================================================================
    doc.add_page_break()
    add_heading_styled(doc, "2. Risk Assessment", level=1)

    doc.add_paragraph(
        "The overall risk grade for this environment is {{ risk_grade }} with a "
        "risk score of {{ risk_score }}. The assessment identified "
        "{{ critical_issues }} critical issues and {{ high_risk_rules }} "
        "high-risk rules requiring immediate attention."
    )

    doc.add_paragraph("")

    risk_table = doc.add_table(rows=4, cols=2)
    risk_table.style = "Table Grid"
    risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    risk_items = [
        ("Best Practices Score", "{{ best_practices_score }}"),
        ("Segmentation Score", "{{ segmentation_score }}"),
        ("Lateral Movement Paths", "{{ lateral_movement_paths }}"),
        ("Shadowed Rules", "{{ shadowed_rule_count }}"),
    ]
    for i, (label, value) in enumerate(risk_items):
        set_cell_shading(risk_table.rows[i].cells[0], "F2F2F2")
        set_cell_text(risk_table.rows[i].cells[0], label, bold=True, size=Pt(9))
        set_cell_text(risk_table.rows[i].cells[1], value, size=Pt(9))

    doc.add_paragraph("")

    # =========================================================================
    # 3. COMPLIANCE
    # =========================================================================
    add_heading_styled(doc, "3. Compliance Summary", level=1)

    doc.add_paragraph(
        "The following table summarises compliance status against applicable "
        "frameworks. This section is populated when compliance data is available "
        "in the Rampart export."
    )

    doc.add_paragraph("")

    comp_table = add_table_with_headers(doc, [
        "Framework", "Score (%)", "Status", "Passed", "Failed", "Total"
    ])
    add_table_marker_row(comp_table, "compliance", [
        "framework", "percentage", "status", "passed", "failed", "total"
    ])

    doc.add_paragraph("")

    # =========================================================================
    # 4. DETAILED FINDINGS
    # =========================================================================
    doc.add_page_break()
    add_heading_styled(doc, "4. Detailed Findings", level=1)

    doc.add_paragraph(
        "This section presents all findings identified during the analysis, "
        "organised by severity."
    )

    doc.add_paragraph("")

    # 4.1 Critical Findings
    add_heading_styled(doc, "4.1 Critical Findings", level=2)
    doc.add_paragraph("{{ critical_count }} critical finding(s) were identified.")
    doc.add_paragraph("")

    crit_table = add_table_with_headers(doc, [
        "Rule", "Device Group", "Severity", "Type", "Description", "Remediation", "Risk"
    ])
    add_table_marker_row(crit_table, "critical_findings", [
        "rule_name", "device_group", "severity", "type", "description", "remediation", "risk_score"
    ])

    doc.add_paragraph("")

    # 4.2 High Findings
    add_heading_styled(doc, "4.2 High Findings", level=2)
    doc.add_paragraph("{{ high_count }} high-severity finding(s) were identified.")
    doc.add_paragraph("")

    high_table = add_table_with_headers(doc, [
        "Rule", "Device Group", "Severity", "Type", "Description", "Remediation", "Risk"
    ])
    add_table_marker_row(high_table, "high_findings", [
        "rule_name", "device_group", "severity", "type", "description", "remediation", "risk_score"
    ])

    doc.add_paragraph("")

    # 4.3 All Findings
    add_heading_styled(doc, "4.3 All Findings", level=2)
    doc.add_paragraph("{{ total_findings }} finding(s) were identified in total.")
    doc.add_paragraph("")

    all_table = add_table_with_headers(doc, [
        "Rule", "Device Group", "Severity", "Type", "Description", "Remediation", "Risk"
    ])
    add_table_marker_row(all_table, "findings", [
        "rule_name", "device_group", "severity", "type", "description", "remediation", "risk_score"
    ])

    doc.add_paragraph("")

    # =========================================================================
    # 5. SHADOWED RULES
    # =========================================================================
    doc.add_page_break()
    add_heading_styled(doc, "5. Shadowed Rules", level=1)

    doc.add_paragraph(
        "{{ shadowed_rules_total }} shadowed rule(s) were detected. Shadowed rules "
        "are never matched because a broader rule higher in the policy takes "
        "precedence."
    )

    doc.add_paragraph("")

    shadow_table = add_table_with_headers(doc, [
        "Rule", "Shadowed By", "Device Group", "Severity", "Description", "Remediation"
    ])
    add_table_marker_row(shadow_table, "shadowed_rules", [
        "rule_name", "shadowed_by", "device_group", "severity", "description", "remediation"
    ])

    doc.add_paragraph("")

    # =========================================================================
    # 6. DUPLICATE OBJECTS
    # =========================================================================
    add_heading_styled(doc, "6. Duplicate Objects", level=1)

    doc.add_paragraph(
        "The analysis identified {{ duplicate_address_count }} duplicate address "
        "object(s) and {{ duplicate_service_count }} duplicate service object(s)."
    )

    doc.add_paragraph("")

    dup_table = add_table_with_headers(doc, [
        "Type", "Value", "Count", "Objects", "Remediation"
    ])
    add_table_marker_row(dup_table, "duplicate_addresses", [
        "type", "value", "count", "objects", "remediation"
    ])

    doc.add_paragraph("")

    # =========================================================================
    # 7. NETWORK SEGMENTATION
    # =========================================================================
    doc.add_page_break()
    add_heading_styled(doc, "7. Network Segmentation", level=1)

    doc.add_paragraph(
        "Segmentation grade: {{ seg_grade }} (score: {{ seg_score }}). "
        "The analysis evaluated {{ seg_zone_count }} zone(s) with "
        "{{ seg_allowed_pairs }} allowed pair(s) and {{ seg_blocked_pairs }} "
        "blocked pair(s)."
    )

    doc.add_paragraph("")

    add_heading_styled(doc, "7.1 Weak Segments", level=2)
    weak_table = add_table_with_headers(doc, [
        "Source Zone", "Destination Zone", "Openness", "Remediation"
    ])
    add_table_marker_row(weak_table, "weak_segments", [
        "source_zone", "dest_zone", "openness", "remediation"
    ])

    doc.add_paragraph("")

    add_heading_styled(doc, "7.2 Lateral Movement", level=2)
    doc.add_paragraph(
        "{{ lateral_movement_count }} rule(s) contribute to lateral movement risk."
    )
    doc.add_paragraph("")

    lat_table = add_table_with_headers(doc, [
        "Rule", "Severity", "Source Zones", "Destination Zones", "Risk Factors"
    ])
    add_table_marker_row(lat_table, "lateral_movement", [
        "rule_name", "severity", "source_zones", "dest_zones", "risk_factors"
    ])

    doc.add_paragraph("")

    # =========================================================================
    # 8. ADDITIONAL ANALYSIS
    # =========================================================================
    doc.add_page_break()
    add_heading_styled(doc, "8. Additional Analysis", level=1)

    # 8.1 Cleartext Protocols
    add_heading_styled(doc, "8.1 Cleartext Protocols", level=2)
    doc.add_paragraph(
        "{{ cleartext_rule_count }} rule(s) permit cleartext protocols."
    )
    doc.add_paragraph("")

    clear_table = add_table_with_headers(doc, [
        "Rule", "Protocol", "Severity", "Secure Alternative"
    ])
    add_table_marker_row(clear_table, "cleartext_rules", [
        "rule_name", "protocol", "severity", "secure_alternative"
    ])

    doc.add_paragraph("")

    # 8.2 Stale Rules
    add_heading_styled(doc, "8.2 Stale Rules", level=2)
    doc.add_paragraph(
        "{{ stale_rule_count }} stale rule(s) were identified."
    )
    doc.add_paragraph("")

    stale_table = add_table_with_headers(doc, [
        "Rule", "Severity", "Indicators", "Disabled"
    ])
    add_table_marker_row(stale_table, "stale_rules", [
        "rule_name", "severity", "indicators", "disabled"
    ])

    doc.add_paragraph("")

    # 8.3 Egress Risks
    add_heading_styled(doc, "8.3 Egress Risks", level=2)
    doc.add_paragraph(
        "{{ egress_risk_count }} rule(s) present egress risk."
    )
    doc.add_paragraph("")

    egress_table = add_table_with_headers(doc, [
        "Rule", "Severity", "Risk Factors", "Remediation"
    ])
    add_table_marker_row(egress_table, "egress_findings", [
        "rule_name", "severity", "risk_factors", "remediation"
    ])

    doc.add_paragraph("")

    # 8.4 Decryption Gaps
    add_heading_styled(doc, "8.4 Decryption Gaps", level=2)
    doc.add_paragraph(
        "{{ decryption_gap_count }} decryption gap(s) were identified."
    )
    doc.add_paragraph("")

    dec_table = add_table_with_headers(doc, [
        "Rule", "Severity", "Reason", "Remediation"
    ])
    add_table_marker_row(dec_table, "decryption_gaps", [
        "rule_name", "severity", "reason", "remediation"
    ])

    doc.add_paragraph("")

    # 8.5 Geo-IP Findings
    add_heading_styled(doc, "8.5 Geo-IP Findings", level=2)
    doc.add_paragraph(
        "{{ geo_ip_unrestricted_count }} rule(s) lack geo-IP restrictions."
    )
    doc.add_paragraph("")

    geo_table = add_table_with_headers(doc, [
        "Rule", "Severity", "Type", "Remediation"
    ])
    add_table_marker_row(geo_table, "geo_ip_findings", [
        "rule_name", "severity", "type", "remediation"
    ])

    doc.add_paragraph("")

    # 8.6 Rule Expiry
    add_heading_styled(doc, "8.6 Rule Expiry", level=2)
    doc.add_paragraph(
        "{{ rule_expiry_count }} rule(s) have expiry concerns."
    )
    doc.add_paragraph("")

    exp_table = add_table_with_headers(doc, [
        "Rule", "Type", "Detail"
    ])
    add_table_marker_row(exp_table, "rule_expiry", [
        "rule_name", "type", "detail"
    ])

    doc.add_paragraph("")

    # =========================================================================
    # 9. RECOMMENDATIONS
    # =========================================================================
    doc.add_page_break()
    add_heading_styled(doc, "9. Recommendations", level=1)

    doc.add_paragraph(
        "Based on the findings of this audit, the following actions are recommended:"
    )

    doc.add_paragraph("")

    rec_items = [
        "Address all {{ critical_count }} critical finding(s) immediately.",
        "Review and remediate {{ high_count }} high-severity finding(s).",
        "Remove or consolidate {{ shadowed_rule_count }} shadowed rule(s) to simplify the policy.",
        "Eliminate {{ duplicate_address_count }} duplicate address object(s) and "
        "{{ duplicate_service_count }} duplicate service object(s).",
        "Review {{ stale_rule_count }} stale rule(s) for decommissioning.",
        "Replace cleartext protocols in {{ cleartext_rule_count }} rule(s) with encrypted alternatives.",
        "Improve network segmentation to reduce lateral movement paths (currently {{ lateral_movement_paths }}).",
    ]

    for item in rec_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    doc.add_paragraph("")

    # =========================================================================
    # APPENDIX
    # =========================================================================
    doc.add_page_break()
    add_heading_styled(doc, "Appendix A: Methodology", level=1)

    doc.add_paragraph(
        "This audit was performed using Rampart, an automated firewall policy "
        "analysis platform. The configuration of type '{{ config_type }}' was "
        "analysed on {{ analysis_timestamp }}."
    )

    doc.add_paragraph("")

    add_heading_styled(doc, "Appendix B: Best Practices", level=1)

    doc.add_paragraph(
        "Best practices assessment score: {{ best_practices_overall_score }} "
        "(grade: {{ best_practices_grade }}). A total of "
        "{{ best_practices_available }} best practice checks were evaluated."
    )

    # -- Save --
    doc.save(output_path)
    print(f"Template created: {output_path}")


if __name__ == "__main__":
    output = sys.argv[1] if len(sys.argv) > 1 else "template.docx"
    create_template(output)
